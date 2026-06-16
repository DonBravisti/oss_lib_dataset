"""
Виртуальная Осетинская Библиотека — построитель датасета (финальная версия)
===========================================================================
Этап 1: рекурсивный обход облака mail.ru → links.csv
Этап 2: текстовый слой + язык → dataset.csv

Детекторы языка:
  LLM:     Mistral, Groq (llama-3.1-8b-instant)
  Модели:  GlotLID (поддерживает осетинский), fasttext-langdetect
  Библ.:   langdetect

Установи зависимости:
    pip install requests pandas mistralai groq python-docx pdfplumber pymupdf
                pdfminer.six ebooklib fasttext-langdetect huggingface_hub
                fasttext-wheel langdetect tqdm

Требования:
    DjVuLibre в PATH (команда djvutxt)

Запуск:
    python build_dataset.py              # полный прогон
    python build_dataset.py --max-files 20  # тест на 20 файлах
"""

import os
import io
import re
import time
import subprocess
import tempfile
import csv
import logging
import argparse
import requests
import pandas as pd
from pathlib import Path
from collections import Counter
from datetime import datetime
from tqdm import tqdm

# ── CONFIG ────────────────────────────────────────────────────────────────────
MISTRAL_API_KEY = "mistral_api_key"
GROQ_API_KEY = "groq_api_key"

CLOUD_WEBLINK = "Cbc3/kLSNPK3A3"
LINKS_CSV = "links.csv"
DATASET_CSV = "dataset.csv"
CSV_SEP = ";"

TEXT_FORMATS = {"djvu", "pdf", "doc", "docx", "fb2",
                "txt", "rtf", "odt", "epub", "htm", "html"}

API_DELAY = 0.25    # пауза между запросами к mail.ru (сек)
LLM_DELAY = 1.5     # пауза между запросами к LLM (сек)
DOWNLOAD_BYTES = 524288  # 512 КБ для анализа
DJVU_PAGES = 6       # страниц DjVu для анализа
MAX_TEXT_LEN = 800     # максимум символов текста для LLM
# ──────────────────────────────────────────────────────────────────────────────


# ══════════════════════════════════════════════════════════════════════════════
# ЛОГИРОВАНИЕ
# ══════════════════════════════════════════════════════════════════════════════

def setup_logging() -> logging.Logger:
    logs_dir = Path("logs")
    logs_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = logs_dir / f"run_{timestamp}.log"

    logger = logging.getLogger("builder")
    logger.setLevel(logging.DEBUG)

    fmt = logging.Formatter(
        "%(asctime)s | %(levelname)-8s | %(message)s", datefmt="%H:%M:%S")

    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)

    ch = logging.StreamHandler()
    ch.setLevel(logging.WARNING)
    ch.setFormatter(fmt)

    logger.addHandler(fh)
    logger.addHandler(ch)

    print(f"  Лог: {log_file}")
    return logger


log = setup_logging()


# ══════════════════════════════════════════════════════════════════════════════
# АРГУМЕНТЫ
# ══════════════════════════════════════════════════════════════════════════════

def parse_args():
    parser = argparse.ArgumentParser(
        description="Построитель датасета Осетинской библиотеки")
    parser.add_argument("--max-files", type=int, default=None,
                        help="Обработать только N файлов (для тестирования)")
    return parser.parse_args()


# ══════════════════════════════════════════════════════════════════════════════
# ИНИЦИАЛИЗАЦИЯ МОДЕЛЕЙ ЯЗЫКА
# ══════════════════════════════════════════════════════════════════════════════

print("Загружаем локальные модели определения языка...")

# fasttext-langdetect
try:
    from ftlangdetect import detect as ft_detect_raw

    def detect_fasttext(text: str) -> str:
        if len(text.strip()) < 20:
            return "unknown"
        try:
            result = ft_detect_raw(text[:1000].replace("\n", " "))
            lang = result.get("lang", "unknown")
            score = result.get("score", 0)
            return lang if score > 0.5 else "unknown"
        except Exception as e:
            log.debug(f"fasttext ошибка: {e}")
            return "unknown"
    print("  ✅ fasttext-langdetect")
except Exception as e:
    log.warning(f"fasttext-langdetect не загружен: {e}")

    def detect_fasttext(text: str) -> str:
        return "unknown"
    print("  ⚠️  fasttext-langdetect недоступен")

# GlotLID — поддерживает осетинский (iron/digor)
try:
    from huggingface_hub import hf_hub_download
    import fasttext
    _glot_path = hf_hub_download(
        repo_id="cis-lmu/glotlid", filename="model.bin", resume_download=True)
    _glot_model = fasttext.load_model(_glot_path)

    def detect_glotlid(text: str) -> str:
        if len(text.strip()) < 20:
            return "unknown"
        try:
            labels, scores = _glot_model.predict(
                text[:1000].replace("\n", " "), k=1)
            lang = labels[0].replace("__label__", "")
            return lang if scores[0] > 0.4 else "unknown"
        except Exception as e:
            log.debug(f"GlotLID ошибка: {e}")
            return "unknown"
    print("  ✅ GlotLID")
except Exception as e:
    log.warning(f"GlotLID не загружен: {e}")

    def detect_glotlid(text: str) -> str:
        return "unknown"
    print("  ⚠️  GlotLID недоступен")

# langdetect
try:
    from langdetect import detect as ld_detect

    def detect_langdetect(text: str) -> str:
        try:
            return ld_detect(text[:500])
        except Exception:
            return "unknown"
    print("  ✅ langdetect")
except Exception as e:
    log.warning(f"langdetect не загружен: {e}")

    def detect_langdetect(text: str) -> str:
        return "unknown"
    print("  ⚠️  langdetect недоступен")


# ══════════════════════════════════════════════════════════════════════════════
# НОРМАЛИЗАЦИЯ И ГОЛОСОВАНИЕ
# ══════════════════════════════════════════════════════════════════════════════

LANG_MAP = {
    # Русский
    "ru": "russian", "rus": "russian", "russian": "russian",
    "rus_cyrl": "russian", "русский": "russian",
    # Осетинский — GlotLID использует коды вида oss_cyrl (iron) и diq (digor)
    "os": "ossetian", "oss": "ossetian", "ossetian": "ossetian",
    "oss_cyrl": "ossetian (iron)", "осетинский": "ossetian",
    "осетинский (иронский диалект)": "ossetian (iron)",
    "осетинский (дигорский диалект)": "ossetian (digor)",
    "diq": "ossetian (digor)", "digor": "ossetian (digor)",
    "iron": "ossetian (iron)",
    "ossetian (iron dialect)": "ossetian (iron)",
    "ossetian (digor dialect)": "ossetian (digor)",
    # Английский
    "en": "english", "eng": "english", "english": "english",
    "английский": "english",
    # Немецкий
    "de": "german", "deu": "german", "german": "german",
    # Французский
    "fr": "french", "fra": "french", "french": "french",
    # Грузинский
    "ka": "georgian", "kat": "georgian", "georgian": "georgian",
    # Армянский
    "hy": "armenian", "hye": "armenian", "armenian": "armenian",
    # Арабский
    "ar": "arabic", "ara": "arabic", "arabic": "arabic",
    # Латинский
    "la": "latin", "lat": "latin", "latin": "latin",
    # Мусор который модели возвращают на плохом тексте
    "zxx": "unknown", "zxx_zzzz": "unknown", "bn": "unknown",
    # Неопределенный язык
    "und_cyrl": "unknown", "und": "unknown",
}


def normalize_lang(raw: str) -> str:
    if not raw:
        return "unknown"
    cleaned = raw.strip().lower()
    if cleaned in LANG_MAP:
        return LANG_MAP[cleaned]
    for key, value in LANG_MAP.items():
        if key in cleaned:
            return value
    return cleaned


SKIP = {"unknown", "no text", "error", "нет текста", "ошибка", ""}

# При ничьей — приоритет этих языков (наиболее ожидаемые в коллекции)
VOTE_PRIORITY = ["russian", "ossetian", "ossetian (iron)",
                 "ossetian (digor)", "english"]


def vote_languages(*langs) -> str:
    normalized = [normalize_lang(l) for l in langs if l not in SKIP]
    if not normalized:
        return "unknown"
    counts = Counter(normalized)
    max_count = counts.most_common(1)[0][1]
    winners = [l for l, c in counts.items() if c == max_count]
    if len(winners) == 1:
        return winners[0]
    for p in VOTE_PRIORITY:
        if p in winners:
            return p
    return winners[0]


# ══════════════════════════════════════════════════════════════════════════════
# СКАЧИВАНИЕ
# ══════════════════════════════════════════════════════════════════════════════

WEBLINK_GET_URL = None


def get_weblink_get_url() -> str:
    global WEBLINK_GET_URL
    if WEBLINK_GET_URL:
        return WEBLINK_GET_URL
    r = requests.get(
        f"https://cloud.mail.ru/api/v2/dispatcher?api=2&_={int(time.time())}",
        headers={"User-Agent": "Mozilla/5.0"}, timeout=15)
    r.raise_for_status()
    WEBLINK_GET_URL = r.json()["body"]["weblink_get"][0]["url"]
    log.info(f"weblink_get: {WEBLINK_GET_URL}")
    return WEBLINK_GET_URL


def download_chunk(weblink: str, num_bytes: int = DOWNLOAD_BYTES) -> bytes | None:
    try:
        url = f"{get_weblink_get_url()}/{weblink}"
        r = requests.get(url,
                         headers={"Range": f"bytes=0-{num_bytes-1}",
                                  "User-Agent": "Mozilla/5.0"},
                         timeout=30)
        if r.status_code in (200, 206):
            log.debug(f"chunk {len(r.content)}B: {weblink}")
            return r.content
        log.warning(f"chunk статус {r.status_code}: {weblink}")
        return None
    except Exception as e:
        log.error(f"chunk ошибка для '{weblink}': {type(e).__name__}: {e}")
        return None


def download_full(weblink: str) -> bytes | None:
    try:
        url = f"{get_weblink_get_url()}/{weblink}"
        r = requests.get(url,
                         headers={"User-Agent": "Mozilla/5.0"}, timeout=120)
        if r.status_code == 200:
            log.debug(f"full {len(r.content)}B: {weblink}")
            return r.content
        log.warning(f"full статус {r.status_code}: {weblink}")
        return None
    except Exception as e:
        log.error(f"full ошибка для '{weblink}': {type(e).__name__}: {e}")
        return None


# ══════════════════════════════════════════════════════════════════════════════
# ИЗВЛЕЧЕНИЕ ТЕКСТА
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf(data: bytes, weblink: str) -> tuple[str, str]:
    # Попытка 1: pymupdf
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        texts = [p.get_text().strip()
                 for p in list(doc)[:6] if p.get_text().strip()]
        text = " ".join(texts)[:MAX_TEXT_LEN]
        if text:
            log.debug(f"PDF pymupdf {len(text)}c: {weblink}")
            return "yes", text
        log.debug(f"PDF pymupdf пустой: {weblink}")
    except Exception as e:
        log.warning(f"PDF pymupdf ошибка '{weblink}': {type(e).__name__}: {e}")

    # Попытка 2: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            texts = [p.extract_text().strip()
                     for p in pdf.pages[:6] if p.extract_text()]
        text = " ".join(texts)[:MAX_TEXT_LEN]
        if text:
            log.debug(f"PDF pdfplumber {len(text)}c: {weblink}")
            return "yes", text
        log.debug(f"PDF pdfplumber пустой: {weblink}")
    except Exception as e:
        log.warning(
            f"PDF pdfplumber ошибка '{weblink}': {type(e).__name__}: {e}")

    # Попытка 3: pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(data), maxpages=6)[:MAX_TEXT_LEN]
        if text.strip():
            log.debug(f"PDF pdfminer {len(text)}c: {weblink}")
            return "yes", text
        log.debug(f"PDF pdfminer пустой: {weblink}")
    except Exception as e:
        log.warning(
            f"PDF pdfminer ошибка '{weblink}': {type(e).__name__}: {e}")

    log.warning(f"PDF: слой есть, текст не извлечён: {weblink}")
    return "yes", ""


def has_text_layer_pdf(data: bytes, weblink: str) -> bool:
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        for page in list(doc)[:5]:
            if len(page.get_text().strip()) > 30:
                return True
        has_fonts = any(page.get_fonts() for page in doc)
        log.debug(f"PDF has_fonts={has_fonts}: {weblink}")
        return has_fonts
    except Exception as e:
        log.warning(
            f"PDF проверка слоя ошибка '{weblink}': {type(e).__name__}: {e}")
        return b"/Font" in data


def extract_djvu(weblink: str) -> tuple[str, str]:
    data = download_full(weblink)
    if data is None:
        return "N/A", ""

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".djvu", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        result = subprocess.run(
            ["djvutxt", f"--page=1-{DJVU_PAGES}", tmp_path],
            capture_output=True, timeout=45,
            encoding="utf-8", errors="replace")

        if result.returncode != 0:
            log.warning(
                f"djvutxt код {result.returncode}: {result.stderr.strip()}: {weblink}")

        text = result.stdout.strip()[:MAX_TEXT_LEN]
        log.debug(f"DjVu djvutxt {len(text)}c: {weblink}")

        if len(text) > 20:
            return "yes", text
        else:
            return "no", ""

    except FileNotFoundError:
        log.error("djvutxt не найден в PATH")
        print("\n    ⚠️  djvutxt не найден — установи DjVuLibre и добавь в PATH")
        return "N/A", ""
    except Exception as e:
        log.error(f"djvutxt ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


def extract_docx(data: bytes, weblink: str) -> tuple[str, str]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = " ".join(p.text for p in doc.paragraphs if p.text.strip())[
            :MAX_TEXT_LEN]
        log.debug(f"DOCX {len(text)}c: {weblink}")
        return ("yes" if len(text.strip()) > 20 else "no"), text
    except Exception as e:
        log.error(f"DOCX ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""


def extract_doc(data: bytes, weblink: str) -> tuple[str, str]:
    tmp_path = None
    # Попытка 1: antiword
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, timeout=30,
            encoding="utf-8", errors="replace")
        text = result.stdout.strip()[:MAX_TEXT_LEN]
        if text:
            log.debug(f"DOC antiword {len(text)}c: {weblink}")
            return "yes", text
        log.debug(f"DOC antiword пустой: {weblink}")
    except FileNotFoundError:
        log.debug(f"DOC: antiword не найден: {weblink}")
    except Exception as e:
        log.warning(
            f"DOC antiword ошибка '{weblink}': {type(e).__name__}: {e}")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            tmp_path = None

    # Попытка 2: python-docx (некоторые .doc это на самом деле docx)
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = " ".join(p.text for p in doc.paragraphs if p.text.strip())[
            :MAX_TEXT_LEN]
        if len(text.strip()) > 20:
            log.debug(f"DOC python-docx {len(text)}c: {weblink}")
            return "yes", text
    except Exception as e:
        log.debug(
            f"DOC python-docx ошибка '{weblink}': {type(e).__name__}: {e}")

    # Попытка 3: cp1251 decode
    try:
        text = data.decode("cp1251", errors="ignore")
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()[:MAX_TEXT_LEN]
        log.debug(f"DOC cp1251 {len(text)}c: {weblink}")
        return ("yes" if len(text.strip()) > 20 else "no"), text
    except Exception as e:
        log.error(f"DOC cp1251 ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""


def extract_epub(data: bytes, weblink: str) -> tuple[str, str]:
    try:
        import ebooklib
        from ebooklib import epub
        book = epub.read_epub(io.BytesIO(data))
        texts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            content = item.get_content().decode("utf-8", errors="replace")
            clean = re.sub(r'<[^>]+>', ' ', content)
            clean = re.sub(r'\s+', ' ', clean).strip()
            if clean:
                texts.append(clean)
            if sum(len(t) for t in texts) > MAX_TEXT_LEN:
                break
        text = " ".join(texts)[:MAX_TEXT_LEN]
        log.debug(f"EPUB {len(text)}c: {weblink}")
        return ("yes" if len(text.strip()) > 20 else "no"), text
    except Exception as e:
        log.error(f"EPUB ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""


def extract_fb2(data: bytes, weblink: str) -> tuple[str, str]:
    try:
        import xml.etree.ElementTree as ET
        for enc in ("utf-8", "cp1251", "windows-1251"):
            try:
                text_data = data.decode(enc)
                break
            except Exception:
                continue
        else:
            log.warning(f"FB2: не удалось декодировать: {weblink}")
            return "N/A", ""
        text_data = re.sub(r' xmlns[^"]*"[^"]*"', '', text_data)
        root = ET.fromstring(text_data)
        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            if sum(len(t) for t in texts) > MAX_TEXT_LEN:
                break
        text = " ".join(texts)[:MAX_TEXT_LEN]
        log.debug(f"FB2 {len(text)}c: {weblink}")
        return ("yes" if len(text.strip()) > 20 else "no"), text
    except Exception as e:
        log.error(f"FB2 ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""


def extract_plaintext(data: bytes, weblink: str) -> tuple[str, str]:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            text = data.decode(enc)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()[:MAX_TEXT_LEN]
            log.debug(f"Plaintext {enc} {len(text)}c: {weblink}")
            return ("yes" if len(text.strip()) > 20 else "no"), text
        except Exception:
            continue
    log.warning(f"Plaintext: не удалось декодировать: {weblink}")
    return "N/A", ""


def get_text_and_layer(weblink: str, fmt: str) -> tuple[str, str]:
    if fmt not in TEXT_FORMATS:
        return "N/A", ""
    try:
        if fmt == "djvu":
            return extract_djvu(weblink)

        data = download_chunk(weblink)
        if data is None:
            return "N/A", ""

        if fmt == "pdf":
            if not has_text_layer_pdf(data, weblink):
                return "no", ""
            return extract_pdf(data, weblink)
        elif fmt == "docx":
            return extract_docx(data, weblink)
        elif fmt == "doc":
            return extract_doc(data, weblink)
        elif fmt == "epub":
            return extract_epub(data, weblink)
        elif fmt == "fb2":
            return extract_fb2(data, weblink)
        elif fmt in ("txt", "htm", "html", "rtf", "odt"):
            return extract_plaintext(data, weblink)
        else:
            return "N/A", ""
    except Exception as e:
        log.error(
            f"get_text_and_layer неожиданная ошибка '{weblink}': {type(e).__name__}: {e}")
        return "N/A", ""


# ══════════════════════════════════════════════════════════════════════════════
# ОПРЕДЕЛЕНИЕ ЯЗЫКА ЧЕРЕЗ LLM
# ══════════════════════════════════════════════════════════════════════════════

LANG_PROMPT = """Detect the language of the following text.
Reply with ONLY the language name in English (e.g.: russian, english, german, ossetian, arabic, etc.).
If the text is in Ossetian, specify the dialect if possible: ossetian (iron) or ossetian (digor).
The text may contain HTML, XML or other markup tags — ignore them and detect the language of the actual text content.
If you cannot determine the language, reply: unknown

Text:
{text}"""


def detect_lang_mistral(text: str) -> str:
    try:
        from mistralai.client import Mistral
        client = Mistral(api_key=MISTRAL_API_KEY)
        resp = client.chat.complete(
            model="mistral-small-latest",
            messages=[{"role": "user",
                       "content": LANG_PROMPT.format(text=text[:500])}],
            max_tokens=30)
        raw = resp.choices[0].message.content.strip()
        result = normalize_lang(raw)
        log.debug(f"Mistral: '{raw}' → '{result}'")
        return result
    except Exception as e:
        log.error(f"Mistral ошибка: {type(e).__name__}: {e}")
        return "error"


def detect_lang_groq(text: str) -> str:
    try:
        from groq import Groq
        client = Groq(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user",
                       "content": LANG_PROMPT.format(text=text[:500])}],
            max_tokens=30)
        raw = resp.choices[0].message.content.strip()
        result = normalize_lang(raw)
        log.debug(f"Groq: '{raw}' → '{result}'")
        return result
    except Exception as e:
        log.error(f"Groq ошибка: {type(e).__name__}: {e}")
        return "error"


NO_TEXT_ROW = {
    "lang_mistral":    "no text",
    "lang_groq":       "no text",
    "lang_fasttext":   "no text",
    "lang_glotlid":    "no text",
    "lang_langdetect": "no text",
    "lang_vote":       "no text",
}


def detect_languages(text: str, filename: str) -> dict:
    if not text or len(text.strip()) < 15:
        log.debug(f"Текст слишком короткий, пропускаем: {filename}")
        return NO_TEXT_ROW

    log.info(f"Определяем язык: {filename}")

    mistral = detect_lang_mistral(text)
    time.sleep(LLM_DELAY)
    groq = detect_lang_groq(text)
    time.sleep(LLM_DELAY)
    fasttext = detect_fasttext(text)
    glotlid = detect_glotlid(text)
    langdetect = detect_langdetect(text)

    vote = vote_languages(mistral, groq, fasttext, glotlid, langdetect)

    log.info(
        f"Результат '{filename}': vote={vote} "
        f"(mistral={mistral}, groq={groq}, "
        f"fasttext={fasttext}, glotlid={glotlid}, langdetect={langdetect})")

    return {
        "lang_mistral":    mistral,
        "lang_groq":       groq,
        "lang_fasttext":   fasttext,
        "lang_glotlid":    glotlid,
        "lang_langdetect": langdetect,
        "lang_vote":       vote,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ЭТАП 1: СБОР ССЫЛОК
# ══════════════════════════════════════════════════════════════════════════════

def make_public_url(weblink: str) -> str:
    encoded = requests.utils.quote(weblink, safe="/")
    return f"https://cloud.mail.ru/public/{encoded}"


def collect_all_files(weblink, rows=None, depth=0):
    if rows is None:
        rows = []
    try:
        r = requests.get(
            "https://cloud.mail.ru/api/v2/folder",
            params={"weblink": weblink, "limit": 500, "offset": 0},
            headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
        data = r.json()["body"]
    except Exception as e:
        log.error(f"Ошибка папки '{weblink}': {e}")
        return rows

    for item in data.get("list", []):
        if item["type"] == "folder":
            print(f"{'  '*depth}📁 {item['name']}")
            collect_all_files(item["weblink"], rows, depth + 1)
        else:
            name = item["name"]
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else "unknown"
            rows.append({
                "filename": name,
                "file_url": make_public_url(item["weblink"]),
                "format":   ext,
                "weblink":  item["weblink"],
            })
            log.debug(f"Файл: {name}")
        time.sleep(API_DELAY)
    return rows


def stage1_collect_links():
    print("=" * 60)
    print("ЭТАП 1: Сбор ссылок")
    print("=" * 60)
    log.info("=== ЭТАП 1 ===")
    rows = collect_all_files(CLOUD_WEBLINK)
    df = pd.DataFrame(rows)
    df.to_csv(LINKS_CSV, index=False, encoding="utf-8-sig", sep=CSV_SEP)
    log.info(f"Этап 1 готов. Файлов: {len(df)}")
    print(f"\n✅ Собрано {len(df)} файлов → {LINKS_CSV}")
    return df


# ══════════════════════════════════════════════════════════════════════════════
# ЭТАП 2: ОСНОВНОЙ ЦИКЛ
# ══════════════════════════════════════════════════════════════════════════════

FIELDNAMES = [
    "filename", "file_url", "format", "has_text_layer",
    "lang_mistral", "lang_groq", "lang_fasttext",
    "lang_glotlid", "lang_langdetect", "lang_vote",
]


def stage2_enrich(df_links: pd.DataFrame, max_files: int = None):
    print("\n" + "=" * 60)
    print("ЭТАП 2: Определение текстового слоя и языка")
    print("=" * 60)
    log.info("=== ЭТАП 2 ===")

    print("Получаем URL для скачивания...")
    get_weblink_get_url()

    # Чекпоинт
    processed_urls = set()
    if Path(DATASET_CSV).exists():
        try:
            df_done = pd.read_csv(
                DATASET_CSV, encoding="utf-8-sig", sep=CSV_SEP)
            processed_urls = set(df_done["file_url"].tolist())
            log.info(f"Чекпоинт: обработано {len(processed_urls)} файлов")
            print(
                f"  Чекпоинт: {len(processed_urls)} файлов уже обработано, продолжаем...")
        except Exception as e:
            log.warning(f"Не удалось прочитать чекпоинт: {e}")

    write_header = not Path(DATASET_CSV).exists()
    out_file = open(DATASET_CSV,
                    "a" if not write_header else "w",
                    encoding="utf-8-sig", newline="")
    writer = csv.DictWriter(out_file, fieldnames=FIELDNAMES, delimiter=CSV_SEP)
    if write_header:
        writer.writeheader()

    # Фильтруем уже обработанные
    to_process = df_links[~df_links["file_url"].isin(processed_urls)]
    if max_files:
        to_process = to_process.head(max_files)
        print(f"  Режим теста: обрабатываем {max_files} файлов")

    total_all = len(df_links)
    for _, row in tqdm(to_process.iterrows(),
                       total=len(to_process), desc="Обработка"):
        idx = row.name
        fmt = str(row["format"]).lower()
        name = row["filename"]
        weblink = str(row["weblink"])
        url = row["file_url"]

        tqdm.write(f"[{idx+1}/{total_all}] {name}", end=" ... ")
        log.info(f"[{idx+1}/{total_all}] {name} | {fmt}")

        has_layer, text_sample = get_text_and_layer(weblink, fmt)
        lang_info = detect_languages(text_sample, name)

        writer.writerow({
            "filename":       name,
            "file_url":       url,
            "format":         fmt,
            "has_text_layer": has_layer,
            **lang_info,
        })
        out_file.flush()
        tqdm.write(f"layer={has_layer}, lang={lang_info['lang_vote']}")
        time.sleep(API_DELAY)

    out_file.close()
    log.info("=== ЭТАП 2 завершён ===")
    print(f"\n✅ Готово! → {DATASET_CSV}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    args = parse_args()
    log.info("=== ЗАПУСК ===")

    if Path(LINKS_CSV).exists():
        print(f"✅ {LINKS_CSV} найден, пропускаем этап 1")
        log.info(f"{LINKS_CSV} найден, пропускаем этап 1")
        df_links = pd.read_csv(LINKS_CSV, encoding="utf-8-sig", sep=CSV_SEP)
    else:
        df_links = stage1_collect_links()

    stage2_enrich(df_links, max_files=args.max_files)
    log.info("=== ЗАВЕРШЕНО ===")
